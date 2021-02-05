import urllib.request
import win32com.client
import docx
import datetime
import requests
from bs4 import BeautifulSoup as bs
import pythoncom

def pars1(name):
    pythoncom.CoInitialize()
    headers = {"accept": "/", "user-agent": 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/76.0.3809.132 Safari/537.36'}
    request = requests.get('http://portal.volgetc.ru/', headers=headers)
    soup = bs(request.content, 'html.parser')
    block_news = soup.findAll('ul', class_="latestnews")[1].find('a')['href']

    request2 = requests.get('http://portal.volgetc.ru/' + str(block_news), headers=headers)
    soup = bs(request2.content, 'html.parser')
    link = soup.findAll('iframe')[1]
    link_twopart = link['src'][33:]

    first_a = link_twopart.split('&')



    file = urllib.request.urlopen(first_a[0]).read()
    f = open("E:\PyCharmProject\TgBot\library\docks/File.doc", "wb")
    f.write(file)
    f.close()

    baseDir = 'E:\PyCharmProject\TgBot\library\docks\File.doc'
    word = win32com.client.Dispatch("Word.application")
    docx_file = '{0}{1}'.format(baseDir, 'x')
    wordDoc = word.Documents.Open(baseDir)
    wordDoc.SaveAs2(docx_file, FileFormat=16)
    wordDoc.Close()
    word.Quit()

    doc = docx.Document("E:\PyCharmProject\TgBot\library\docks\File.docx")
    paragraphs = doc.paragraphs[1]
    base = []

    for table in doc.tables:
        amount_rows = len(table.rows)
        i=0
        while i < amount_rows:
            section = table.cell(i, 0)
            if name == str(section.text):
                num_lesson = table.cell(i, 1)
                subject = table.cell(i, 2)
                teacher = table.cell(i, 3)
                num_class = table.cell(i, 4)

                base.append(section.text + ' ' + num_lesson.text + ' ' + subject.text + ' ' + teacher.text + ' ' + num_class.text)
            i+=1

    main_str = str(paragraphs.text) + '\n'
    for i in range(len(base)):
        main_str = main_str + base[i] + '\n'
    print(main_str)
    return main_str



