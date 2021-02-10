import urllib.request
import win32com.client
import pythoncom
from bs4 import BeautifulSoup as bs
import requests
import docx
from library import bd



def connect_site():

    headers = {"accept": "/",
                "user-agent": 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/76.0.3809.132 Safari/537.36'}
    request = requests.get('http://portal.volgetc.ru/', headers=headers)
    soup = bs(request.content, 'html.parser')
    block_news = soup.findAll('ul', class_="latestnews")[1].find('a')['href']

    request2 = requests.get('http://portal.volgetc.ru/' + str(block_news), headers=headers)
    soup = bs(request2.content, 'html.parser')
    link = soup.findAll('iframe')[1]
    link_twopart = link['src'][33:]

    global first_a
    first_a = link_twopart.split('&')

def doc_work():
    pythoncom.CoInitialize()
    file = urllib.request.urlopen(first_a[0]).read()
    f = open("E:\PyCharmProject\TgBot\library\docks/File.doc", "wb")
    f.write(file)
    f.close()

    baseDir = 'E:\PyCharmProject\TgBot\library\docks\File.doc'
    word = win32com.client.Dispatch("Word.application")
    docx_file = '{0}{1}'.format(baseDir, 'x')
    wordDoc = word.Documents.Open(baseDir)
    wordDoc.SaveAs2(docx_file, FileFormat=16)
    wordDoc.Close()
    word.Quit()

def pars():

    connect_site()
    doc_work()

    doc = docx.Document("E:\PyCharmProject\TgBot\library\docks\File.docx")

    paragraphs1 = ''
    for i in range(0,3):
        paragraphs1 = paragraphs1 + doc.paragraphs[i].text


    #Подключени к бд
    conn = bd.connect_bd()

    query = ("""DELETE FROM timetable""")
    bd.update_bd(query, conn)


    id = 0
    for table in doc.tables:
        amount_rows = len(table.rows)
        i = 0
        print(len(doc.tables))

        while i < amount_rows:
            try:
                id+=1

                section = table.cell(i, 0)
                num_lesson = table.cell(i, 1)
                subject = table.cell(i, 2)
                teacher = table.cell(i, 3)
                num_class = table.cell(i, 4)

                query = ("""INSERT INTO timetable (id, name_group, pair_number, subject, teacher_fullname, room_number, date) 
	                            VALUES 
		                            ({0},'{1}','{2}', '{3}', '{4}', '{5}', '{6}')
                             """).format(id,section.text,num_lesson.text,subject.text,teacher.text, num_class.text, paragraphs1)
                bd.update_bd(query,conn)

            except:
                pass
            i += 1
    conn.close()


